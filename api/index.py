from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import anthropic
import os
import zipfile
import tempfile
import pandas as pd
import docx
import PyPDF2
import json
import re
from io import BytesIO
import openpyxl
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Get API key from environment variables
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")

# Initialize Anthropic client
client = anthropic.Anthropic(
    api_key=ANTHROPIC_API_KEY,
)

# Function to extract content from Excel files
def extract_excel_content(file_path):
    workbook = openpyxl.load_workbook(file_path)
    content = []
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        df = pd.DataFrame(sheet.values)
        content.append(f"Sheet: {sheet_name}\n{df.to_string(index=False, header=False)}")
    return "\n\n".join(content)

# Function to extract content from PDF files
def extract_pdf_content(file_path):
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        content = []
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            content.append(f"Page {page_num + 1}:\n{page.extract_text()}")
        return "\n\n".join(content)

# Function to recursively process directory and extract file contents
def process_directory(directory):
    content = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            file_path = os.path.join(root, file)
            rel_path = os.path.relpath(file_path, directory)
            
            if file.endswith('.csv'):
                try:
                    df = pd.read_csv(file_path)
                    content.append(f"CSV content of {rel_path}:\n{df.to_string()}")
                except Exception as e:
                    content.append(f"Error reading CSV {rel_path}: {str(e)}")
            
            elif file.endswith('.txt'):
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content.append(f"Text content of {rel_path}:\n{f.read()}")
                except Exception as e:
                    content.append(f"Error reading text file {rel_path}: {str(e)}")
            
            elif file.endswith('.docx'):
                try:
                    doc = docx.Document(file_path)
                    text = '\n'.join([para.text for para in doc.paragraphs])
                    content.append(f"DOCX content of {rel_path}:\n{text}")
                except Exception as e:
                    content.append(f"Error reading DOCX {rel_path}: {str(e)}")
            
            elif file.endswith('.md'):
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content.append(f"Markdown content of {rel_path}:\n{f.read()}")
                except Exception as e:
                    content.append(f"Error reading Markdown {rel_path}: {str(e)}")
            
            elif file.endswith(('.xls', '.xlsx')):
                try:
                    content.append(f"Excel content of {rel_path}:\n{extract_excel_content(file_path)}")
                except Exception as e:
                    content.append(f"Error reading Excel {rel_path}: {str(e)}")
            
            elif file.endswith('.pdf'):
                try:
                    content.append(f"PDF content of {rel_path}:\n{extract_pdf_content(file_path)}")
                except Exception as e:
                    content.append(f"Error reading PDF {rel_path}: {str(e)}")
            
            elif file.endswith('.json'):
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        json_content = json.load(f)
                        content.append(f"JSON content of {rel_path}:\n{json.dumps(json_content, indent=2)}")
                except Exception as e:
                    content.append(f"Error reading JSON {rel_path}: {str(e)}")
            
            else:
                # Try to read as text for unknown file types
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content.append(f"Content of {rel_path} (unknown format):\n{f.read()}")
                except Exception:
                    content.append(f"Unable to read file {rel_path}")
    
    return '\n\n'.join(content)

# Function to extract content from different file types
async def extract_file_content(file: UploadFile):
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_file_path = os.path.join(temp_dir, file.filename)
        with open(temp_file_path, 'wb') as f:
            content = await file.read()
            f.write(content)
        
        if file.filename.endswith('.zip'):
            # Handle zip files
            extraction_dir = os.path.join(temp_dir, "extracted")
            os.makedirs(extraction_dir, exist_ok=True)
            
            try:
                # Extract the zip file
                with zipfile.ZipFile(temp_file_path, 'r') as zip_ref:
                    zip_ref.extractall(extraction_dir)
                
                # Process the extracted contents
                return process_directory(extraction_dir)
            except zipfile.BadZipFile:
                return "Invalid ZIP file format"
        
        elif file.filename.endswith('.csv'):
            try:
                df = pd.read_csv(temp_file_path)
                return f"CSV content:\n{df.to_string()}"
            except Exception as e:
                return f"Error reading CSV: {str(e)}"
        
        elif file.filename.endswith('.txt'):
            with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f"Text content:\n{f.read()}"
        
        elif file.filename.endswith('.docx'):
            doc = docx.Document(temp_file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            return f"DOCX content:\n{text}"
        
        elif file.filename.endswith('.md'):
            with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f"Markdown content:\n{f.read()}"
        
        elif file.filename.endswith(('.xls', '.xlsx')):
            return f"Excel content:\n{extract_excel_content(temp_file_path)}"
        
        elif file.filename.endswith('.pdf'):
            return f"PDF content:\n{extract_pdf_content(temp_file_path)}"
        
        elif file.filename.endswith('.json'):
            with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                json_content = json.load(f)
                return f"JSON content:\n{json.dumps(json_content, indent=2)}"
        
        else:
            # Try to read as text for unknown file types
            try:
                with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f"Content of file (unknown format):\n{f.read()}"
            except Exception:
                return f"Unable to read file {file.filename}"

@app.post("/api/")
async def process_request(question: str = Form(...), file: UploadFile = None):
    try:
        file_content = ""
        if file:
            file_content = await extract_file_content(file)
        
        # Create a prompt with the question and file content
        prompt = f"""You are an AI assistant helping a student with their IIT Madras Online Degree in Data Science assignment questions.
The student needs the exact answer to enter in their graded assignment.

Question: {question}

File Content:
{file_content}

Instructions:
1. Provide ONLY the exact answer to the question, without any explanations or extra text.
2. The answer should be ready to be directly entered into the assignment form.
3. If the question asks for values from a CSV file or data, extract those specific values.
4. If the answer is a number, provide just the number.
5. If the answer is text, provide just the text.
6. Do not include any explanations, citations, or your thought process.

Your answer should be extremely concise and exactly match what is required for the assignment.
"""
        
        # Use Claude to generate the answer
        print(prompt)
        message = client.messages.create(
            model="claude-3-7-sonnet-20250219",
            max_tokens=20000,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": prompt
                        }
                    ]
                }
            ],
            thinking={
                "type": "enabled",
                "budget_tokens": 16000
            }
        )
        print(message)
        
        # Extract the answer from Claude's response and clean it
        answer = ""
        for item in message.content:
            if item["type"] == "text":
                answer = item["text"]
        # if message.content and len(message.content) > 0:
        #     first_content = message.content[0]
        #     if first_content.type == "text":
        #         answer = first_content.text.strip()
        
        # Remove any markdown or code block formatting from the answer
        # answer = re.sub(r'```.*?\n', '', answer)  # Remove code block start
        # answer = re.sub(r'```', '', answer)       # Remove code block end
        # answer = re.sub(r'^Answer: ', '', answer) # Remove "Answer:" prefix
        # answer = answer.strip()
        
        # Return the answer in the required format
        return {"answer": answer}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")

@app.get("/api/health")
async def health_check():
    return {"status": "ok", "message": "API is running"}

# Handler for serverless deployment on Vercel
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
